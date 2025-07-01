import streamlit as st
import fitz                 # PyMuPDF
import docx                 # python-docx
from fpdf import FPDF       # fpdf2
import google.generativeai as genai
import textwrap, io, pathlib
import os
from datetime import datetime
import re
from docx import Document
from io import BytesIO

# ─── Configuration ────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Resume Enhancer",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ─── API Setup ────────────────────────────────────────────────────────────────
def setup_gemini():
    """Setup Gemini API with proper error handling"""
    api_key = None
    
    # Try multiple sources for API key
    try:
        api_key = st.secrets["GEMINI_API_KEY"]
    except (KeyError, FileNotFoundError, Exception):
        api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        st.error("🔑 **API Key Required**")
        st.markdown("""
        Please provide your Google Gemini API key. You can:
        1. Set it as an environment variable: `GEMINI_API_KEY=your_key_here`
        2. Create a `.streamlit/secrets.toml` file with: `GEMINI_API_KEY = \"your_key_here\"`
        3. Enter it in the sidebar (not recommended for production)
        
        Get your API key from: https://makersuite.google.com/app/apikey
        """)
        with st.sidebar:
            st.warning("⚠️ Development Mode")
            api_key = st.text_input(
                "Enter Gemini API Key",
                type="password",
                help="This is for development only. Use environment variables or secrets.toml for production."
            )
        if not api_key:
            st.stop()
    try:
        genai.configure(api_key=api_key)  # type: ignore[reportPrivateImportUsage]
        model = genai.GenerativeModel("gemini-1.5-flash")  # type: ignore[reportPrivateImportUsage]
        return model
    except Exception as e:
        st.error(f"Failed to initialize Gemini: {str(e)}")
        st.stop()

# ─── Helper Functions ─────────────────────────────────────────────────────────
def extract_text_from_pdf(uploaded_file) -> str:
    """Extract text from PDF file with error handling"""
    try:
        # Use explicit page loading to avoid linter error
        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        text = "\n".join(doc.load_page(i).get_text("text") for i in range(doc.page_count))  # type: ignore
        doc.close()
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return ""

def extract_text_from_docx(uploaded_file) -> str:
    """Extract text from DOCX file with error handling"""
    try:
        doc = docx.Document(uploaded_file)
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)
    except Exception as e:
        st.error(f"Error reading DOCX: {str(e)}")
        return ""

def clean_text(text: str) -> str:
    """Clean and normalize extracted text"""
    # Remove excessive whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = re.sub(r' +', ' ', text)
    return text.strip()

def build_analysis_prompt(resume_text: str) -> str:
    """Build comprehensive analysis prompt"""
    return f"""
You are a professional resume analyst and career coach. Analyze the following resume comprehensively:

RESUME TEXT:
{resume_text}

ANALYSIS REQUIREMENTS:
1. Overall Score (0-100) with justification
2. Strengths (what works well)
3. Weaknesses (specific areas needing improvement)
4. ATS Compatibility assessment
5. Industry-specific recommendations
6. Missing Information Needed (list specific questions to ask the candidate)

Format your response clearly with headers for each section.
"""

def build_rewrite_prompt(resume_text: str, extra_info: str, style_preferences: dict) -> str:
    """Build rewrite prompt with style preferences and instruct AI to fill in gaps intelligently"""
    style_guide = f"""
STYLE PREFERENCES:
- Format: {style_preferences['format']}
- Tone: {style_preferences['tone']}
- Length: {style_preferences['length']}
- Industry Focus: {style_preferences['industry']}
"""
    
    return f"""
You are an expert resume writer. Using the information below, create a complete, professional, and ATS-friendly resume.
- Use all available details from the original resume and the user's answers.
- If any section is missing or incomplete, intelligently fill in the gaps with realistic, relevant content based on the candidate's background and the target industry.
- Improve clarity, phrasing, and formatting throughout.
- Do NOT include placeholders or incomplete sections.
- Return ONLY the finished resume, no commentary.

{style_guide}

EXISTING RESUME:
{resume_text}

ADDITIONAL INFORMATION FROM USER:
{extra_info}

REQUIREMENTS:
- Use strong action verbs
- Quantify achievements where possible
- Ensure proper formatting and structure
- Include relevant keywords for ATS

RESUME:
"""

def create_pdf_from_text(resume_text: str, filename: str = "resume") -> bytes:
    """Create PDF with better formatting and error handling"""
    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        # Add DejaVuSans font (Unicode)
        pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
        pdf.add_font("DejaVu", "B", "DejaVuSans-Bold.ttf", uni=True)
        pdf.set_font("DejaVu", size=11)
        lines = resume_text.split('\n')
        usable_width = pdf.w - 2 * pdf.l_margin  # Safe width for multicell

        for line in lines:
            line = line.strip()
            if line:
                if line.isupper() or line.endswith(':'):
                    pdf.set_font("DejaVu", style='B', size=12)
                    pdf.ln(2)
                    # Wrap long lines, break long words
                    wrapped = textwrap.fill(line, width=90, break_long_words=True)
                    pdf.multi_cell(usable_width, 6, wrapped)
                    pdf.ln(1)
                else:
                    pdf.set_font("DejaVu", size=10)
                    wrapped = textwrap.fill(line, width=90, break_long_words=True)
                    pdf.multi_cell(usable_width, 5, wrapped)
            else:
                pdf.ln(3)
        # Output PDF as bytes using dest='S'. Handle both str and bytearray return types.
        pdf_data = pdf.output(dest='S')
        if isinstance(pdf_data, str):
            pdf_bytes = pdf_data.encode('latin1')
        else:
            pdf_bytes = bytes(pdf_data)
        return pdf_bytes
    except Exception as e:
        st.error(f"Error creating PDF: {str(e)}")
        return b""

def create_docx_from_text(resume_text: str, filename: str = "resume") -> bytes:
    """Create a DOCX file from resume text and return as bytes."""
    doc = Document()
    for line in resume_text.split('\n'):
        if line.strip() == "":
            doc.add_paragraph()  # Blank line
        else:
            doc.add_paragraph(line)
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.read()

def extract_questions_from_analysis(analysis_text: str) -> list:
    """Extract questions from analysis with improved parsing"""
    questions = []
    lines = analysis_text.split('\n')
    
    # Look for common question indicators
    question_indicators = [
        "missing information", "information needed", "questions",
        "additional details", "clarification needed"
    ]
    
    in_question_section = False
    for line in lines:
        line_lower = line.lower().strip()
        
        # Check if we're entering a question section
        if any(indicator in line_lower for indicator in question_indicators):
            in_question_section = True
            continue
        
        # Stop at next major section
        if in_question_section and line.strip() and not line.startswith((' ', '-', '•', '1', '2', '3', '4', '5')):
            if ':' in line and len(line.split(':')[0].split()) < 4:
                in_question_section = False
        
        # Extract questions
        if in_question_section and line.strip():
            # Clean up question text
            question = re.sub(r'^[-•\d\.\s]+', '', line).strip()
            if question and len(question) > 10:  # Filter out too short lines
                questions.append(question)
    
    # Fallback questions if none found
    if not questions:
        questions = [
            "What specific achievements or metrics can you provide?",
            "Are there any additional skills or certifications to highlight?",
            "What is your target job role or industry?",
            "Any additional experiences or projects to include?"
        ]
    
    return questions[:6]  # Limit to 6 questions

# ─── Main Application ─────────────────────────────────────────────────────────
def main():
    st.title("📄 AI-Powered Resume Enhancer")
    st.markdown("Transform your resume with AI analysis and optimization")
    
    # Initialize Gemini
    model = setup_gemini()
    
    # Sidebar for settings
    with st.sidebar:
        st.header("⚙️ Settings")
        
        # Style preferences
        st.subheader("Resume Style")
        format_style = st.selectbox(
            "Format Style",
            ["Professional", "Modern", "Creative", "Academic"],
            help="Choose the overall formatting style"
        )
        
        tone_style = st.selectbox(
            "Tone",
            ["Professional", "Dynamic", "Conservative", "Technical"],
            help="Select the writing tone"
        )
        
        length_pref = st.selectbox(
            "Length Preference",
            ["Concise (1 page)", "Standard (1-2 pages)", "Detailed (2+ pages)"],
            help="Preferred resume length"
        )
        
        industry_focus = st.text_input(
            "Industry Focus",
            placeholder="e.g., Technology, Finance, Healthcare",
            help="Target industry for optimization"
        )
        
        style_preferences = {
            'format': format_style,
            'tone': tone_style,
            'length': length_pref,
            'industry': industry_focus or "General"
        }
    
    # File upload
    uploaded_files = st.file_uploader(
        "📁 Upload your resume files",
        type=["pdf", "docx"],
        accept_multiple_files=True,
        help="Upload PDF or DOCX files for analysis"
    )
    
    if not uploaded_files:
        st.info("👆 Upload your resume file(s) to get started")
        return
    
    # Process each uploaded file
    for idx, uploaded_file in enumerate(uploaded_files):
        with st.container():
            st.header(f"📄 {uploaded_file.name}")
            
            # Create columns for better layout
            col1, col2 = st.columns([2, 1])
            
            with col1:
                # Extract text
                with st.spinner("Extracting text..."):
                    if uploaded_file.type == "application/pdf":
                        resume_text = extract_text_from_pdf(uploaded_file)
                    else:
                        resume_text = extract_text_from_docx(uploaded_file)
                    
                    resume_text = clean_text(resume_text)
                
                if not resume_text:
                    st.error("Could not extract text from file")
                    continue
                
                # Show preview
                with st.expander("📖 Text Preview", expanded=False):
                    st.text_area(
                        "Extracted Text",
                        resume_text[:2000] + ("..." if len(resume_text) > 2000 else ""),
                        height=200,
                        disabled=True
                    )
            
            with col2:
                st.metric("Characters", len(resume_text))
                st.metric("Words", len(resume_text.split()))
                st.metric("Lines", len(resume_text.split('\n')))
            
            # Analysis section
            analysis_key = f"analysis_{idx}_{uploaded_file.name}"
            
            if st.button(f"🔍 Analyze Resume", key=f"analyze_{idx}", type="primary"):
                with st.spinner("🤖 AI is analyzing your resume..."):
                    try:
                        analysis_response = model.generate_content(build_analysis_prompt(resume_text))
                        st.session_state[analysis_key] = analysis_response.text
                    except Exception as e:
                        st.error(f"Analysis failed: {str(e)}")
            
            # Show analysis if available
            if analysis_key in st.session_state:
                analysis = st.session_state[analysis_key]
                
                with st.expander("🧠 AI Analysis Results", expanded=True):
                    st.markdown(analysis)
                
                # Extract and display questions
                questions = extract_questions_from_analysis(analysis)
                
                # Enhancement form
                with st.form(key=f"enhancement_form_{idx}"):
                    st.subheader("✏️ Provide Additional Information")
                    st.markdown("*Help the AI create a better resume by answering these questions:*")
                    
                    answers = {}
                    for i, question in enumerate(questions):
                        answer = st.text_area(
                            f"**Q{i+1}:** {question}",
                            key=f"answer_{idx}_{i}",
                            height=80,
                            placeholder="Your answer here..."
                        )
                        answers[i] = answer
                    
                    # Submit button
                    submitted = st.form_submit_button(
                        "🚀 Generate Enhanced Resume",
                        type="primary",
                        use_container_width=True
                    )
                
                # Generate enhanced resume
                if submitted:
                    # Process answers
                    processed_answers = []
                    for i, question in enumerate(questions):
                        answer = answers.get(i, "").strip()
                        if answer:
                            processed_answers.append(f"Q: {question}\nA: {answer}")
                    
                    if not processed_answers:
                        st.warning("Please provide at least one answer to improve your resume.")
                    else:
                        extra_info = "\n\n".join(processed_answers)
                        
                        with st.spinner("🎯 Creating your enhanced resume..."):
                            try:
                                if model is None:
                                    st.error("Gemini API client not initialized. Check your API key.")
                                else:
                                    enhanced_response = model.generate_content(build_rewrite_prompt(resume_text, extra_info, style_preferences))
                                    enhanced_resume = enhanced_response.text or ""
                                
                                # Create DOCX
                                docx_bytes = create_docx_from_text(enhanced_resume)
                                
                                if docx_bytes:
                                    # Success message and download
                                    st.success("✅ Enhanced resume created successfully!")
                                    
                                    # Create filename with timestamp
                                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                    filename = f"Enhanced_Resume_{timestamp}.docx"
                                    
                                    col1, col2 = st.columns(2)
                                    with col1:
                                        st.download_button(
                                            "📥 Download Enhanced Resume (DOCX)",
                                            data=docx_bytes,
                                            file_name=filename,
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            type="primary",
                                            use_container_width=True
                                        )
                                    
                                    with col2:
                                        if st.button("👀 Preview Enhanced Resume", key=f"preview_{idx}"):
                                            st.session_state[f"show_preview_{idx}"] = True
                                    
                                    # Show preview if requested
                                    if st.session_state.get(f"show_preview_{idx}", False):
                                        with st.expander("📄 Enhanced Resume Preview", expanded=True):
                                            st.text_area(
                                                "Enhanced Resume",
                                                enhanced_resume,
                                                height=400,
                                                disabled=True
                                            )
                                else:
                                    st.error("Failed to create DOCX. Please try again.")
                                    
                            except Exception as e:
                                st.error(f"Enhancement failed: {str(e)}")
            
            st.divider()

if __name__ == "__main__":
    main()